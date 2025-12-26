"""
Word Document Report Generator for Metal Loss Assessment
"""
import io
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading_with_style(doc: Document, text: str, level: int = 1):
    """Add a styled heading to the document."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_parameter_table(doc: Document, inputs: Dict):
    """Add input parameters table to document."""
    add_heading_with_style(doc, "Input Parameters", level=2)
    
    # Create table data
    table_data = [
        ("Feature Number/Identifier", str(inputs.get('feature_ID', ''))),
        ("Depth of the defect, %NWT", f"{inputs['dimp_org_percent']}"),
        ("Length of the defect, mm", f"{inputs['Limp_org']}"),
        ("Outside diameter of the pipe, mm", f"{inputs['do']}"),
        ("Pipe wall thickness, mm", f"{inputs['tp']}"),
        ("Yield strength of the material, MPa", f"{inputs['YS']}"),
        ("Tensile strength of the material, MPa", f"{inputs['TS']}"),
        ("ILI vendor", inputs.get('vendor_ILI', '')),
        ("ILI time", inputs.get('date_ILI', '')),
        ("ILI tool tolerance for depth, %", f"{inputs['ILI_dimp_tolerance']}"),
        ("ILI tool tolerance for length, mm", f"{inputs['ILI_Limp_tolerance']}"),
        ("Feature depth growth rate, mm/yr", f"{inputs['CR_low']} (low), {inputs['CR_ave']} (avg), {inputs['CR_high']} (high)"),
        ("Feature length growth rate, mm/yr", f"{inputs.get('CR_Limp', 0)}"),
    ]
    
    # Create table
    table = doc.add_table(rows=len(table_data) + 1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    header_cells[0].text = "Parameter, unit"
    header_cells[1].text = "Value"
    
    # Format headers
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Add data
    for i, (param, value) in enumerate(table_data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = param
        row_cells[1].text = value
    
    doc.add_paragraph()  # Add spacing


def add_data_table(doc: Document, title: str, df: pd.DataFrame, caption: str = ""):
    """Add a data table to the document."""
    add_heading_with_style(doc, title, level=2)
    
    if caption:
        p = doc.add_paragraph(caption)
        p.style = 'Intense Quote'
    
    # Create table
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        header_cells[i].text = str(col)
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Add data
    for i, row in df.iterrows():
        row_cells = table.rows[i + 1].cells
        for j, value in enumerate(row):
            if pd.isna(value):
                row_cells[j].text = ""
            elif isinstance(value, float):
                row_cells[j].text = f"{value:.2f}"
            else:
                row_cells[j].text = str(value)
    
    doc.add_paragraph()  # Add spacing


def add_chart_image(doc: Document, image_path: str, title: str, width: float = 6.0):
    """Add a chart image to the document."""
    add_heading_with_style(doc, title, level=2)
    
    try:
        doc.add_picture(image_path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"Error adding chart: {str(e)}")
    
    doc.add_paragraph()  # Add spacing


def generate_word_report(
    assessment_results: Dict,
    chart_images: Dict[str, bytes],
    output_path: str = None
) -> bytes:
    """
    Generate a Word document report for metal loss assessment.
    
    Parameters:
    -----------
    assessment_results : dict
        Complete assessment results from assess_metal_loss_feature()
    chart_images : dict
        Dictionary with keys: 'depth_growth', 'sop_decay', 'sop_cutoff'
        Values are PNG image bytes
    output_path : str, optional
        Path to save document. If None, returns bytes
    
    Returns:
    --------
    bytes
        Word document as bytes
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('Metal Loss Feature Assessment Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Description
    add_heading_with_style(doc, "Description", level=1)
    doc.add_paragraph(
        "The PNG Asset Integrity team conducted an assessment of a metal loss feature. "
        "This report presents the evaluation results based on modified B31G methodology."
    )
    doc.add_paragraph()
    
    # Inputs and Assumptions
    add_heading_with_style(doc, "Inputs and Assumptions", level=1)
    doc.add_paragraph(
        "The inputs used to evaluate the anomaly were taken from the PNG Integrity Dig Package. "
        "Three different conditions were evaluated:"
    )
    
    inputs = assessment_results['inputs']
    bullets = [
        f"50th Percentile Corrosion Growth Rate of {inputs['CR_low']} mm/yr",
        f"90th Percentile Corrosion Growth Rate of {inputs['CR_ave']} mm/yr",
        f"99th Percentile Corrosion Growth Rate of {inputs['CR_high']} mm/yr"
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        "The selected corrosion rates for this evaluation are based on industry standards "
        "and Bayesian evaluation using in-line inspection data. The evaluation was completed "
        "with in-house developed Python code using modified B31G to calculate the failure "
        "pressure of a given feature. The Safe Operation Pressure (SOP) is calculated using "
        "failure pressure divided by a safety factor of 1.25."
    )
    doc.add_paragraph()
    
    # Input parameters table
    add_parameter_table(doc, inputs)
    
    doc.add_page_break()
    
    # Depth Growth Chart
    if 'depth_growth' in chart_images and chart_images['depth_growth']:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
            tmp.write(chart_images['depth_growth'])
            tmp_path = tmp.name
        add_chart_image(doc, tmp_path, "Evaluation Results - Depth Growth Figure")
        Path(tmp_path).unlink()
    else:
        doc.add_paragraph("⚠️ Depth Growth Chart could not be generated due to image export issues.")
    
    # Depth Growth Table
    months = list(range(inputs['month_CR']))
    depth_low = assessment_results['depth_arrays']['low']
    depth_ave = assessment_results['depth_arrays']['ave']
    depth_high = assessment_results['depth_arrays']['high']
    tp = inputs['tp']
    
    df_depth = pd.DataFrame({
        'Month': months,
        'Depth Low (mm)': depth_low,
        'Depth Low (%)': [d / tp * 100 for d in depth_low],
        'Depth Ave (mm)': depth_ave,
        'Depth Ave (%)': [d / tp * 100 for d in depth_ave],
        'Depth High (mm)': depth_high,
        'Depth High (%)': [d / tp * 100 for d in depth_high]
    })
    
    add_data_table(
        doc,
        "Evaluation Results - Depth Growth Table",
        df_depth.head(20),  # Show first 20 rows
        caption="Table: Predicted Defect Depth Growth Over Time (first 20 months shown)"
    )
    
    doc.add_page_break()
    
    # SOP Decay Chart
    if 'sop_decay' in chart_images and chart_images['sop_decay']:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
            tmp.write(chart_images['sop_decay'])
            tmp_path = tmp.name
        add_chart_image(doc, tmp_path, "Evaluation Results - SOP Decay Figure")
        Path(tmp_path).unlink()
    else:
        doc.add_paragraph("⚠️ SOP Decay Chart could not be generated due to image export issues.")
    
    # SOP Decay Table
    sop_low = assessment_results['sop_arrays']['low']
    sop_ave = assessment_results['sop_arrays']['ave']
    sop_high = assessment_results['sop_arrays']['high']
    
    df_sop = pd.DataFrame({
        'Month': months,
        'SOP Low (psi)': sop_low,
        'SOP Ave (psi)': sop_ave,
        'SOP High (psi)': sop_high
    })
    
    add_data_table(
        doc,
        "Evaluation Results - SOP Decay Table",
        df_sop.head(20),  # Show first 20 rows
        caption="Table: Predicted Safe Operating Pressure Decay Over Time (first 20 months shown)"
    )
    
    doc.add_page_break()
    
    # SOP Decay with Cutoff Chart
    if 'sop_cutoff' in chart_images and chart_images['sop_cutoff']:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
            tmp.write(chart_images['sop_cutoff'])
            tmp_path = tmp.name
        add_chart_image(doc, tmp_path, "Evaluation Results - SOP Decay with 80% Wall Thickness Cutoff")
        Path(tmp_path).unlink()
    else:
        doc.add_paragraph("⚠️ SOP Decay with Cutoff Chart could not be generated due to image export issues.")
    
    # Cutoff information
    add_heading_with_style(doc, "80% Wall Thickness Cutoff Information", level=2)
    cutoff_months = assessment_results['cutoff_months']
    month_CR = assessment_results['inputs']['month_CR']  # Extract month_CR from inputs
    for rate_name, rate_key in [("Low", "low"), ("Average", "ave"), ("High", "high")]:
        cutoff = cutoff_months[rate_key]
        if cutoff == -1:
            doc.add_paragraph(f"{rate_name} corrosion rate: Safe for full projection period ({month_CR} months)")
        else:
            doc.add_paragraph(f"{rate_name} corrosion rate: Reaches 80% wall thickness at Month {cutoff}")
    
    # Save or return bytes
    if output_path:
        doc.save(output_path)
        with open(output_path, 'rb') as f:
            return f.read()
    else:
        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.read()

